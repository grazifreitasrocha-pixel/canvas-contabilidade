import os, re, shutil, urllib.parse, requests, socket, logging
import asyncio, time, sys
import json as _json_lib
import subprocess
from datetime import datetime, timedelta
from pathlib import Path
from dotenv import load_dotenv
from flask import Flask, request, jsonify
import threading
from telegram.error import Conflict
# Log em arquivo para diagnóstico
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler(Path(__file__).parent / "canvas_bot.log", encoding="utf-8"),
        logging.StreamHandler(),
    ],
)
logger = logging.getLogger(__name__)

# Força IPv4 para evitar falhas de conexão via IPv6
_orig_getaddrinfo = socket.getaddrinfo
def _getaddrinfo_ipv4(host, port, _family=0, *args, **kwargs):
    return _orig_getaddrinfo(host, port, socket.AF_INET, *args, **kwargs)
socket.getaddrinfo = _getaddrinfo_ipv4
from docx import Document
from telegram import Update, ReplyKeyboardMarkup, ReplyKeyboardRemove
from telegram.ext import (Application, CommandHandler, ConversationHandler,
                           MessageHandler, filters, ContextTypes)

load_dotenv()

BOT_TOKEN  = os.environ["BOT_TOKEN"]
GROUP_ID   = int(os.environ["TELEGRAM_GROUP_ID"])
TRELLO_KEY = os.environ["TRELLO_KEY"]
TRELLO_TOK = os.environ["TRELLO_API_TOKEN"]
ASAAS_KEY  = os.environ.get("ASAAS_KEY", "")
D4SIGN_TOKEN   = os.environ.get("D4SIGN_TOKEN_API", "")
D4SIGN_CRYPT   = os.environ.get("D4SIGN_CRYPT_KEY", "")
D4SIGN_SAFE    = os.environ.get("D4SIGN_SAFE_UUID", "")
D4SIGN_API_URL = "https://secure.d4sign.com.br/api/v1"

PENDENTES_FILE = Path(__file__).parent / "pendentes_assinatura.json"

# Card modelo com checklist de integração
TEMPLATE_CARD_INTEGRACAO = "KbA4dhyW"

# IDs das listas de entrada em cada quadro
ENTRY_LISTS = {
    "1": "694197174aa205b99b63d0be",  # Integração → 1. Venda Ganha
    "2": "69a2f7279d5270c9fdc3b434",  # Alvarás    → 🔵 Novas Solicitações
    "3": "69a594121db830f4646b9518",  # Serviços   → 1. ENTRADA DE SOLICITAÇÃO
    "7": "694197174aa205b99b63d0be",  # MEI        → Integração (mesma lista)
}
TIPOS = {
    "1": "Integração do Cliente",
    "2": "Alvarás e Licenças",
    "3": "Abertura de empresa",
    "4": "Certificado Digital",
    "5": "Serviços da Folha",
    "6": "Serviços do Fiscal",
    "7": "MEI",
}

# Tipos que NÃO têm Trello
SEM_TRELLO = {"4", "5", "6"}

def converter_docx_para_pdf(caminho_docx: Path) -> Path:
    pasta_saida = caminho_docx.parent
    subprocess.run(
        [
            "libreoffice", "--headless", "--convert-to", "pdf",
            "--outdir", str(pasta_saida), str(caminho_docx)
        ],
        check=True,
        capture_output=True,
    )
    caminho_pdf = pasta_saida / (caminho_docx.stem + ".pdf")
    if not caminho_pdf.exists():
        raise FileNotFoundError(f"PDF nao foi gerado: {caminho_pdf}")
    return caminho_pdf


def enviar_para_d4sign(caminho_pdf: Path, email_cliente: str,
                       nome_cliente: str) -> str:
    import time as _time
    headers = {
        "Content-Type": "application/json",
        "Accept": "application/json",
    }
    params_auth = {
        "tokenAPI": D4SIGN_TOKEN,
        "cryptKey": D4SIGN_CRYPT,
    }

    with open(caminho_pdf, "rb") as f:
        resp = requests.post(
            f"{D4SIGN_API_URL}/documents/{D4SIGN_SAFE}/upload",
            params=params_auth,
           files={"file": (caminho_pdf.name, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
    resp.raise_for_status()
    uuid_doc = resp.json().get("uuid")
    if not uuid_doc:
        raise ValueError(f"D4Sign nao retornou UUID: {resp.text}")

    _time.sleep(3)

    resp = requests.post(
        f"{D4SIGN_API_URL}/documents/{uuid_doc}/createlist",
        params=params_auth,
        headers=headers,
        json=[{
            "email": email_cliente,
            "act": "1",
            "foreign": "0",
            "certificadoicpbr": "0",
            "assinatura_presencial": "0",
            "embed_methodauth": "email",
        }],
    )
    resp.raise_for_status()

    resp = requests.post(
        f"{D4SIGN_API_URL}/documents/{uuid_doc}/sendtosign",
        params=params_auth,
        headers=headers,
        json={
            "message": (
                f"Olá, {nome_cliente}! "
                "Segue o contrato de prestação de serviços da Canvas Contabilidade "
                "para sua assinatura eletrônica. Qualquer dúvida, entre em contato conosco."
            ),
            "skip_email": "0",
            "workflow": "0",
            "tokenAPI": D4SIGN_TOKEN,
        },
    )
    resp.raise_for_status()

    webhook_url = f"{URL_BASE_RAILWAY}/webhook/d4sign"
    requests.post(
        f"{D4SIGN_API_URL}/documents/{uuid_doc}/webhooks",
        params=params_auth,
        headers=headers,
        json={"url": webhook_url},
    )

    return uuid_doc


def salvar_pendente(uuid_doc: str, dados: dict) -> None:
    pendentes = {}
    if PENDENTES_FILE.exists():
        try:
            pendentes = _json_lib.loads(PENDENTES_FILE.read_text(encoding="utf-8"))
        except Exception:
            pendentes = {}
    pendentes[uuid_doc] = dados
    PENDENTES_FILE.write_text(
        _json_lib.dumps(pendentes, ensure_ascii=False, indent=2),
        encoding="utf-8"
    )


def ler_e_remover_pendente(uuid_doc: str) -> dict | None:
    if not PENDENTES_FILE.exists():
        return None
    try:
        pendentes = _json_lib.loads(PENDENTES_FILE.read_text(encoding="utf-8"))
    except Exception:
        return None
    dados = pendentes.pop(uuid_doc, None)
    PENDENTES_FILE.write_text(
        _json_lib.dumps(pendentes, ensure_ascii=False, indent=2),
        encoding="utf-8"
    )
    return dados

def trello_add_card(tipo: str, nome: str, desc: str) -> str:
    """
    Cria um cartao novo no Trello, na lista de entrada correspondente
    ao tipo de servico, e copia os checklists do cartao modelo
    (TEMPLATE_CARD_INTEGRACAO) para o cartao recem-criado.

    Retorna a URL do cartao criado.
    """
    list_id = ENTRY_LISTS.get(tipo)
    if not list_id:
        raise ValueError(f"Tipo '{tipo}' nao tem lista de entrada configurada")

    # Cria o cartao
    r = requests.post(
        "https://api.trello.com/1/cards",
        params={"key": TRELLO_KEY, "token": TRELLO_TOK},
        json={"name": nome, "desc": desc, "idList": list_id, "pos": "top"},
    )
    r.raise_for_status()
    card = r.json()
    card_id = card["id"]
    card_url = card.get("shortUrl") or card.get("url")

    # Copia os checklists do cartao modelo, se houver um configurado
    # para este tipo de servico (so faz sentido para Integracao, "1")
    if tipo == "1" and TEMPLATE_CARD_INTEGRACAO:
        try:
            r_checklists = requests.get(
                f"https://api.trello.com/1/cards/{TEMPLATE_CARD_INTEGRACAO}/checklists",
                params={"key": TRELLO_KEY, "token": TRELLO_TOK},
            )
            r_checklists.raise_for_status()
            for checklist in r_checklists.json():
                requests.post(
                    "https://api.trello.com/1/checklists",
                    params={"key": TRELLO_KEY, "token": TRELLO_TOK},
                    json={"idCard": card_id, "idChecklistSource": checklist["id"]},
                ).raise_for_status()
        except Exception:
            # Se copiar o checklist falhar, o cartao ja foi criado -
            # nao queremos que isso quebre o cadastro inteiro
            pass

    return card_url

def calcular_competencia(data_texto: str) -> str:
    """
    Recebe a data de inicio digitada (ex: '25/06/2026') e retorna
    a competencia no formato 'Mes/Ano' (ex: 'Junho/2026').
    """
    texto = data_texto.strip()
    partes = re.split(r"[/\-]", texto)

    if len(partes) != 3:
        return "verificar manualmente"

    try:
        mes_num = int(partes[1])
        ano_num = int(partes[2])
    except (ValueError, IndexError):
        return "verificar manualmente"

    if not (1 <= mes_num <= 12):
        return "verificar manualmente"

    if ano_num < 100:
        ano_num += 2000

    nome_mes = MESES_PT[mes_num - 1].capitalize()
    return f"{nome_mes}/{ano_num}"

CONTRACT_TMPL     = Path(__file__).parent / "templates" / "Contrato Assessoria.docx.docx"
CONTRACT_TMPL_MEI = Path(__file__).parent / "templates" / "Contrato_MEI_Canvas_Revisado.docx"
CONTRATOS_DIR = Path(__file__).parent / "contratos"
ASAAS_URL     = "https://api.asaas.com/v3/payments"

# ── Estados ────────────────────────────────────────────────────────────────────
(TIPO, EMPRESA, CNPJ, REGIME, DATA_INICIO, SERVICOS, FOLHA,
 HONORARIO, VENCIMENTO, RESPONSAVEL, OBS, REPRESENTANTE,
 CPF_REP, END_EMPRESA, END_REP, WHATSAPP, EMAIL) = range(17)

# Estados do fluxo /boleto
(BOL_EMPRESA, BOL_CNPJ, BOL_DESC, BOL_VALOR, BOL_VENC, BOL_WHATSAPP) = range(17, 23)



# ── Asaas ──────────────────────────────────────────────────────────────────────

def _asaas(method: str, path: str, **kwargs):
    headers = {"access_token": ASAAS_KEY, "Content-Type": "application/json"}
    r = requests.request(method, f"{ASAAS_URL}{path}", headers=headers, **kwargs)
    if not r.ok:
        try:
            erros = r.json().get("errors", [])
            detail = "; ".join(e.get("description", str(e)) for e in erros) if erros else r.text
        except Exception:
            detail = r.text
        raise Exception(f"Asaas {r.status_code}: {detail}")
    return r.json()

def asaas_customer(name: str, cpf_cnpj: str) -> str:
    clean = re.sub(r"\D", "", cpf_cnpj)
    found = _asaas("GET", "/customers", params={"cpfCnpj": clean}).get("data", [])
    if found:
        return found[0]["id"]
    return _asaas("POST", "/customers", json={"name": name, "cpfCnpj": clean})["id"]

def asaas_boleto(customer_id: str, value: float, due_day: int, desc: str) -> dict:
    import time
    now = datetime.now()
    m   = now.month + (1 if now.day >= due_day else 0)
    y   = now.year + (1 if m > 12 else 0)
    m   = m % 12 or 12
    due = datetime(y, m, due_day).strftime("%Y-%m-%d")
    d   = _asaas("POST", "/payments", json={
        "customer": customer_id, "billingType": "BOLETO",
        "value": value, "dueDate": due, "description": desc,
    })
    payment_id = d["id"]
    url = d.get("bankSlipUrl") or d.get("invoiceUrl", "")
    if not url:
        time.sleep(2)
        info = _asaas("GET", f"/payments/{payment_id}/bankSlipUrl")
        url  = info.get("bankSlipUrl", "")
    return {"id": payment_id, "url": url, "due": due}


# ── Contrato ───────────────────────────────────────────────────────────────────

MESES_PT = ["janeiro","fevereiro","março","abril","maio","junho",
            "julho","agosto","setembro","outubro","novembro","dezembro"]

def _apply_subst(doc, subst: dict, bracket_style: bool):
    def sub_para(para):
        full = "".join(r.text for r in para.runs)
        new  = full
        if bracket_style:
            for key, val in subst.items():
                new = new.replace(key, val)
        else:
            for key, val in subst.items():
                new = re.sub(r'\$\{"' + re.escape(key) + r'"\s*:\s*"[^"]*"\}', val, new)
        if new != full and para.runs:
            para.runs[0].text = new
            for run in para.runs[1:]:
                run.text = ""

    for para in doc.paragraphs:
        sub_para(para)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    sub_para(para)

def fill_contract_mei(d: dict, today, safe: str) -> Path:
    doc   = Document(CONTRACT_TMPL_MEI)
    subst = {
        "[Nome completo do Cliente]":              d["representante"],
        "[CPF ou CNPJ do Contratante]":            d["cpf_rep"],
        "[CNPJ do MEI]":                           d["cnpj"],
        "[Endereço completo com CEP]":             d["end_empresa"],
        "[E-mail do Contratante]":                 d.get("email", "-"),
        "[Telefone]":                              d["whatsapp"],
        "[Valor da mensalidade (ex: R$ 89,90)]":  f"R$ {d['honorario']}",
        "[valor por extenso]":                     d["honorario"],
        "[dia do vencimento (ex: 10)]":            d["vencimento"],
        "[data de início (dd/mm/aaaa)]":           d["data_inicio"],
        "[dia]":                                   today.strftime("%d"),
        "[mês]":                                   MESES_PT[today.month - 1],
        "[ano]":                                   today.strftime("%Y"),
        "[NOME COMPLETO DO CONTRATANTE]":          d["representante"].upper(),
    }
    _apply_subst(doc, subst, bracket_style=True)
    out = CONTRATOS_DIR / f"Contrato_MEI_{safe}_{today.strftime('%Y%m%d_%H%M')}.docx"
    doc.save(out)
    return out

def fill_contract_assessoria(d: dict, today, safe: str) -> Path:
    doc   = Document(CONTRACT_TMPL)
    subst = {
        "CONTRATANTE":                  d["empresa"],
        "CNPJ":                         d["cnpj"],
        "Endereço da empresa":          d["end_empresa"],
        "Representante legal":          d["representante"],
        "CPF":                          d["cpf_rep"],
        "Endereço contratante":         d["end_rep"],
        "Início contrato":              d["data_inicio"],
        "Valor do honorário":           f"R$ {d['honorario']}",
        "Data de vencimento":           f"dia {d['vencimento']}",
        "Data e local de assinatura":   f"Fortaleza, {today.strftime('%d/%m/%Y')}",
        "Contratante":                  d["empresa"],
    }
    _apply_subst(doc, subst, bracket_style=False)
    out = CONTRATOS_DIR / f"Contrato_{safe}_{today.strftime('%Y%m%d_%H%M')}.docx"
    doc.save(out)
    return out

def fill_contract(d: dict) -> Path:
    CONTRATOS_DIR.mkdir(exist_ok=True)
    today = datetime.now()
    safe  = re.sub(r"[^\w\s-]", "", d["empresa"])[:40].strip()
    if d["tipo"] == "7":
        return fill_contract_mei(d, today, safe)
    return fill_contract_assessoria(d, today, safe)


# ── Mensagem grupo ─────────────────────────────────────────────────────────────
REGUA_PRAZOS = [
    {"dias_min": 0, "dias_max": 5, "responsavel": "Bruna", "area": "Cadastral"},
    {"dias_min": 6, "dias_max": 10, "responsavel": "Fernanda", "area": "Fiscal"},
    {"dias_min": 11, "dias_max": 15, "responsavel": "Daniela", "area": "Folha"},
]


def calcular_prazos(data_assinatura_texto: str) -> list:
    """
    Recebe a data de assinatura (ex: '25/06/2026') e devolve uma lista
    de dicionarios com a data-limite calculada para cada responsavel,
    baseado na REGUA_PRAZOS (dias corridos a partir da assinatura).
    """
    texto = data_assinatura_texto.strip()
    partes = re.split(r"[/\-]", texto)

    if len(partes) != 3:
        return []

    try:
        dia, mes, ano = int(partes[0]), int(partes[1]), int(partes[2])
        if ano < 100:
            ano += 2000
        data_base = datetime(ano, mes, dia)
    except (ValueError, IndexError):
        return []

    resultado = []
    for regra in REGUA_PRAZOS:
        data_limite = data_base + timedelta(days=regra["dias_max"])
        resultado.append({
            "responsavel": regra["responsavel"],
            "area": regra["area"],
            "data_limite": data_limite.strftime("%d/%m/%Y"),
        })
    return resultado


# Certifique-se de ter este import no topo do arquivo (junto com
# "from datetime import datetime"):
#
#   from datetime import datetime, timedelta

def build_group_msg(d: dict, trello_url: str) -> str:
    label = TIPOS[d["tipo"]]
    competencia = calcular_competencia(d["data_inicio"])
    prazos = calcular_prazos(d["data_inicio"])

    modalidade = d.get("modalidade", "Transferência")
    cnpj_exibir = d["cnpj"] if d["cnpj"] else "Aguardando emissão"

    # Deriva "Comercio" a partir do texto em servicos
    palavras_comercio = ["comércio", "comercio", "venda", "produtos", "mercadoria", "loja", "varejo"]
    texto_servicos = d.get("servicos", "").lower()
    comercio = "Sim" if any(p in texto_servicos for p in palavras_comercio) else "Não"

    # Monta o bloco de prazos e responsaveis
    bloco_prazos = ""
    if prazos:
        linhas_prazo = []
        for p in prazos:
            # So mostra a linha de Folha se o cliente realmente tem folha
            if p["area"] == "Folha" and d.get("folha", "").strip().lower() != "sim":
                continue
            linhas_prazo.append(f"• {p['area']} — até {p['data_limite']} ({p['responsavel']})")
        if linhas_prazo:
            bloco_prazos = (
                f"{'─'*30}\n"
                f"*📤 Prazos e responsáveis:*\n"
                + "\n".join(linhas_prazo) + "\n"
            )

    # Próxima etapa do fluxo, baseada no tipo de servico
    proxima_etapa = "Bruna (Cadastral)" if d["tipo"] not in SEM_TRELLO else "—"

    return (
        f"🆕 NOVO CADASTRO\n"
        f"🔖 Tipo: {label}\n"
        f"🔁 *Modalidade: {modalidade.upper()}*\n"
        f"{'─'*30}\n\n"
        f"🏢 Empresa: {d['empresa']}\n"
        f"📄 CNPJ/CPF: {cnpj_exibir}\n"
        f"📱 WhatsApp: {d['whatsapp']}\n"
        f"💼 Regime tributário: {d['regime']}\n"
        f"🏪 Comércio: {comercio}\n"
        f"{'─'*30}\n"
        f"📅 Data de assinatura: {d['data_inicio']}\n"
        f"🗓️ *COMPETÊNCIA: {competencia.upper()}*\n"
        f"{'─'*30}\n"
        f"📋 Serviços contratados: {d['servicos']}\n"
        f"👥 Folha de pagamento: {d['folha']}\n"
        f"💰 Honorário: R$ {d['honorario']} | Venc. {d['vencimento']}\n"
        f"{bloco_prazos}"
        f"{'─'*30}\n"
        f"📌 Status atual: Novo cadastro\n"
        f"➡️ *Próxima etapa: {proxima_etapa.upper()}*\n"
        f"👩‍💼 Cadastrado por: {d['responsavel']}\n"
        f"{'─'*30}\n"
        f"📝 Obs: {d['obs']}\n"
        f"🔗 Ver cartão no Trello: {trello_url}\n\n"
        f"Canvas Contabilidade — Aqui a gente não para"
    )

# ── Handlers da conversa ───────────────────────────────────────────────────────

async def novo(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data.clear()
    await update.message.reply_text(
        "🆕 *NOVO CADASTRO*\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n"
        "Qual o tipo de serviço?\n\n"
        "1️⃣ — Integração do Cliente\n"
        "2️⃣ — Alvarás e Licenças\n"
        "3️⃣ — Abertura / Alteração / Baixa de Empresa\n"
        "4️⃣ — Certificado Digital\n"
        "5️⃣ — Serviços da Folha\n"
        "6️⃣ — Serviços do Fiscal\n"
        "7️⃣ — MEI\n\n"
        "Digite o número:",
        parse_mode="Markdown",
        reply_markup=ReplyKeyboardMarkup(
            [["1"], ["2"], ["3"], ["4"], ["5"], ["6"], ["7"]],
            one_time_keyboard=True, resize_keyboard=True
        )
    )
    return TIPO

async def get_tipo(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    t = update.message.text.strip()[0]
    if t not in TIPOS:
        await update.message.reply_text("Digite 1, 2 ou 3.")
        return TIPO
    ctx.user_data["tipo"] = t
    await update.message.reply_text("🏢 Nome da empresa:", reply_markup=ReplyKeyboardRemove())
    return EMPRESA

async def get_empresa(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["empresa"] = update.message.text.strip().upper()
    await update.message.reply_text("📄 CNPJ ou CPF:")
    return CNPJ

async def get_cnpj(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["cnpj"] = update.message.text.strip()
    await update.message.reply_text(
        "💼 Regime tributário:",
        reply_markup=ReplyKeyboardMarkup(
            [["MEI"], ["Simples Nacional"], ["Lucro Presumido"], ["Lucro Real"]],
            one_time_keyboard=True, resize_keyboard=True
        )
    )
    return REGIME

async def get_regime(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["regime"] = update.message.text.strip()
    await update.message.reply_text("📅 Início da responsabilidade (dd/mm/aaaa):", reply_markup=ReplyKeyboardRemove())
    return DATA_INICIO

async def get_data(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["data_inicio"] = update.message.text.strip()
    await update.message.reply_text("📋 Serviços contratados:")
    return SERVICOS

async def get_servicos(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["servicos"] = update.message.text.strip()
    await update.message.reply_text(
        "👥 Folha de pagamento?",
        reply_markup=ReplyKeyboardMarkup([["Sim"], ["Não"]], one_time_keyboard=True, resize_keyboard=True)
    )
    return FOLHA

async def get_folha(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["folha"] = update.message.text.strip()
    await update.message.reply_text("💰 Honorário mensal (ex: 200,00):", reply_markup=ReplyKeyboardRemove())
    return HONORARIO

async def get_honorario(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["honorario"] = update.message.text.strip()
    await update.message.reply_text("📅 Dia de vencimento (ex: 15):")
    return VENCIMENTO

async def get_vencimento(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["vencimento"] = update.message.text.strip()
    await update.message.reply_text("👩‍💼 Responsável pela integração:")
    return RESPONSAVEL

async def get_responsavel(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["responsavel"] = update.message.text.strip()
    await update.message.reply_text("📝 Observações (ou 'nenhuma'):")
    return OBS

async def get_obs(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["obs"] = update.message.text.strip()
    await update.message.reply_text("👤 Representante legal (nome completo):")
    return REPRESENTANTE

async def get_representante(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["representante"] = update.message.text.strip()
    await update.message.reply_text("🪪 CPF do representante:")
    return CPF_REP

async def get_cpf_rep(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["cpf_rep"] = update.message.text.strip()
    await update.message.reply_text("📍 Endereço da empresa (completo):")
    return END_EMPRESA

async def get_end_empresa(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["end_empresa"] = update.message.text.strip()
    await update.message.reply_text("🏠 Endereço do representante (completo):")
    return END_REP

async def get_end_rep(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["end_rep"] = update.message.text.strip()
    await update.message.reply_text("📱 WhatsApp do cliente (com DDD, ex: 85999887766):")
    return WHATSAPP

async def get_whatsapp(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["whatsapp"] = re.sub(r"\D", "", update.message.text.strip())
    await update.message.reply_text("📧 E-mail do cliente:")
    return EMAIL

async def get_email(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["email"] = update.message.text.strip()
    d = ctx.user_data

    await update.message.reply_text("⏳ Processando...", reply_markup=ReplyKeyboardRemove())

    erros      = []
    trello_url = "—"
    boleto_url = "—"

    # 1 ── Trello (apenas para tipos com quadro)
    if d["tipo"] not in SEM_TRELLO:
        try:
            desc = (
                f"**Empresa:** {d['empresa']}\n"
                f"**CNPJ/CPF:** {d['cnpj']}\n"
                f"**Regime:** {d['regime']}\n"
                f"**Inicio:** {d['data_inicio']}\n"
                f"**Servicos:** {d['servicos']}\n"
                f"**Folha:** {d['folha']}\n"
                f"**Honorario:** R$ {d['honorario']} | Venc. {d['vencimento']}\n"
                f"**Responsavel:** {d['responsavel']}\n"
                f"**Obs:** {d['obs']}"
            )
            trello_url = trello_add_card(
                d["tipo"],
                f"{d['empresa']} — {TIPOS[d['tipo']]}",
                desc
            )
        except Exception as e:
            erros.append(f"Trello: {e}")
    else:
        trello_url = "sem Trello para este servico"

    # 2 ── Mensagem no grupo Telegram
    try:
        await ctx.bot.send_message(chat_id=GROUP_ID, text=build_group_msg(d, trello_url))
    except Exception as e:
        erros.append(f"Grupo Telegram (chat_id={GROUP_ID}): {e}")

    # 3 ── Contrato
    contract_path = None
    try:
        contract_path = fill_contract(d)
    except Exception as e:
        erros.append(f"Contrato: {e}")

    # 4 ── Boleto Asaas
    if ASAAS_KEY and ASAAS_KEY != "PREENCHER_COM_CHAVE_ASAAS":
        try:
            valor = float(d["honorario"].replace(".", "").replace(",", "."))
            cid   = asaas_customer(d["empresa"], d["cnpj"])
            bol   = asaas_boleto(cid, valor, int(d["vencimento"]),
                                 f"Honorario {TIPOS[d['tipo']]} — {d['empresa']}")
            boleto_url = bol["url"]
        except Exception as e:
            erros.append(f"Asaas: {e}")
    else:
        erros.append("Asaas: chave nao configurada — boleto nao gerado")

    # 5 ── Informa caminho do contrato (sem upload para evitar timeout)
    if contract_path and contract_path.exists():
        erros_info = f"\n\nContrato salvo em:\n{contract_path}"
    else:
        erros_info = ""

    # 6 ── Link WhatsApp
    if d["tipo"] == "1":
        wa_text = (
            f"🎉 Seja muito bem-vindo(a) à Canvas Contabilidade!\n\n"
            f"Eu sou a Bruna, da equipe de implantação, e estarei com você durante todo o processo inicial da sua empresa conosco. 😊\n\n"
            f"Nosso objetivo é tornar essa integração simples, rápida e tranquila, garantindo que você tenha segurança, clareza e um atendimento próximo desde o primeiro dia.\n\n"
            f"Sempre que surgir qualquer dúvida, pode contar comigo! 🤝\n\n"
            f"⏰ Nosso horário de atendimento:\n"
            f"Segunda a sexta-feira\n"
            f"07h30 às 11h00 e das 13h00 às 17h30\n\n"
            f"Para darmos andamento à integração da sua empresa na Canvas Contabilidade, precisamos dos seguintes documentos:\n\n"
            f"👥 Dos sócios:\n"
            f"• Documento pessoal (RG e CPF ou CNH)\n"
            f"• Comprovante de endereço atualizado\n"
            f"• Certidão de casamento (caso seja casado)\n\n"
            f"🏢 Da empresa:\n"
            f"• Contrato social e alterações (preferencialmente em PDF)\n"
            f"• Cartão do CNPJ\n"
            f"• Comprovante de endereço atualizado\n"
            f"• Contrato de locação do imóvel (caso seja alugado)\n\n"
            f"Assim que recebermos a documentação, iniciaremos as próximas etapas da implantação. ✨\n\n"
            f"Será um prazer fazer parte dessa nova fase da sua empresa!\n\n"
            f"Link de pagamento do boleto de honorário:\n{boleto_url}"
        )
    elif d["tipo"] == "3":
        wa_text = (
            f"Olá, {d['empresa'].title()}! Tudo bem?\n\n"
            f"É com muita alegria que compartilhamos com você o Contrato Social e o Cartão CNPJ da sua empresa. 🎉\n\n"
            f"A partir de agora, você inicia uma nova fase como empresário, e queremos que saiba que estamos aqui para caminhar ao seu lado em cada passo.\n\n"
            f"Nos próximos dias, o nosso time de implantação entrará em contato para orientar sobre os primeiros procedimentos, obrigações e também para realizar um treinamento, garantindo que você tenha clareza e segurança nessa jornada.\n\n"
            f"Link de pagamento do boleto de honorário:\n{boleto_url}\n\n"
            f"Um grande abraço,\n\n"
            f"Equipe Canvas Contabilidade"
        )
    else:
        wa_text = (
            f"Olá, {d['empresa'].title()}! 😊\n\n"
            f"Seja bem-vindo(a) à Canvas Contabilidade!\n\n"
            f"Serviço contratado: {TIPOS[d['tipo']]}\n\n"
            f"Segue o link do seu boleto:\n{boleto_url}\n\n"
            f"Vencimento: dia {d['vencimento']}\n"
            f"Valor: R$ {d['honorario']}\n\n"
            f"Qualquer dúvida, estamos à disposição! 🟢\n"
            f"Canvas Contabilidade"
        )
    wa_link = f"https://wa.me/55{d['whatsapp']}?text={urllib.parse.quote(wa_text)}"

    # Resumo final
    status = "✅ Tudo certo!" if not erros else "⚠️ Concluído com avisos:\n" + "\n".join(f"• {e}" for e in erros)

    await update.message.reply_text(
        f"{status}\n\n"
        f"Trello: {trello_url}\n"
        f"Boleto: {boleto_url}\n\n"
        f"WhatsApp:\n{wa_link}"
        f"{erros_info}",
        disable_web_page_preview=True
    )

    return ConversationHandler.END

# ── Fluxo /boleto ──────────────────────────────────────────────────────────────

async def boleto_start(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data.clear()
    await update.message.reply_text(
        "💰 *Boleto de Honorário Avulso*\n\n🏢 Nome da empresa ou cliente:",
        parse_mode="Markdown",
        reply_markup=ReplyKeyboardRemove(),
    )
    return BOL_EMPRESA

async def bol_empresa(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["bol_empresa"] = update.message.text.strip().upper()
    await update.message.reply_text("📄 CNPJ ou CPF:")
    return BOL_CNPJ

async def bol_cnpj(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["bol_cnpj"] = update.message.text.strip()
    await update.message.reply_text("📋 Descrição do serviço (ex: Honorário maio/2026):")
    return BOL_DESC

async def bol_desc(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["bol_desc"] = update.message.text.strip()
    await update.message.reply_text("💰 Valor (ex: 350,00):")
    return BOL_VALOR

async def bol_valor(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["bol_valor"] = update.message.text.strip()
    await update.message.reply_text("📅 Dia de vencimento (ex: 10):")
    return BOL_VENC

async def bol_venc(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["bol_venc"] = update.message.text.strip()
    await update.message.reply_text("📱 WhatsApp do cliente (com DDD, ex: 85999887766):")
    return BOL_WHATSAPP

async def bol_whatsapp(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["bol_whatsapp"] = re.sub(r"\D", "", update.message.text.strip())
    d = ctx.user_data

    await update.message.reply_text("⏳ Gerando boleto...")

    try:
        valor = float(d["bol_valor"].replace(".", "").replace(",", "."))
        cid   = asaas_customer(d["bol_empresa"], d["bol_cnpj"])
        bol   = asaas_boleto(cid, valor, int(d["bol_venc"]), d["bol_desc"])

        wa_text = (
            f"Olá, {d['bol_empresa'].title()}! 😊\n\n"
            f"Segue o boleto referente a:\n{d['bol_desc']}\n\n"
            f"💰 Valor: R$ {d['bol_valor']}\n"
            f"📅 Vencimento: {bol['due'].split('-')[2]}/{bol['due'].split('-')[1]}/{bol['due'].split('-')[0]}\n\n"
            f"🔗 {bol['url']}\n\n"
            f"Qualquer dúvida, estamos à disposição! 🟢\n"
            f"Canvas Contabilidade"
        )
        wa_link = f"https://wa.me/55{d['bol_whatsapp']}?text={urllib.parse.quote(wa_text)}"

        await update.message.reply_text(
            f"✅ *Boleto gerado!*\n\n"
            f"🏢 {d['bol_empresa']}\n"
            f"📋 {d['bol_desc']}\n"
            f"💰 R$ {d['bol_valor']}\n"
            f"📅 Venc: {bol['due'].split('-')[2]}/{bol['due'].split('-')[1]}/{bol['due'].split('-')[0]}\n\n"
            f"🔗 Boleto: {bol['url']}\n\n"
            f"📲 Enviar pelo WhatsApp:\n{wa_link}",
            parse_mode="Markdown",
            disable_web_page_preview=True,
        )
    except Exception as e:
        await update.message.reply_text(f"❌ Erro ao gerar boleto: {e}")

    ctx.user_data.clear()
    return ConversationHandler.END


async def cancelar(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("❌ Cancelado.", reply_markup=ReplyKeyboardRemove())
    return ConversationHandler.END

async def start(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "Canvas Contabilidade 🟢\n\n"
        "/novo — cadastrar novo cliente\n"
        "/boleto — gerar boleto de honorário avulso\n"
        "/cancelar — cancelar operação em andamento"
    )

async def chatid(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(f"Chat ID: {update.effective_chat.id}")

async def error_handler(update: object, ctx: ContextTypes.DEFAULT_TYPE):
    logger.error("Erro no handler", exc_info=ctx.error)
    if isinstance(update, Update) and update.effective_message:
        try:
            await update.effective_message.reply_text(
                "⚠️ Ocorreu um erro inesperado. Use /cancelar e tente novamente."
            )
        except Exception:
            pass

# ── Main ───────────────────────────────────────────────────────────────────────
async def meuid(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    chat_id = update.message.chat.id
    chat_type = update.message.chat.type
    await update.message.reply_text(f"ID deste chat: {chat_id}\nTipo: {chat_type}")

def build_app():
    from telegram.ext import Application as _App
    app = _App.builder().token(BOT_TOKEN).build()

    conv = ConversationHandler(
        entry_points=[CommandHandler("novo", novo)],
        name="novo",
        persistent=False,
        states={
            TIPO:         [MessageHandler(filters.TEXT & ~filters.COMMAND, get_tipo)],
            EMPRESA:      [MessageHandler(filters.TEXT & ~filters.COMMAND, get_empresa)],
            CNPJ:         [MessageHandler(filters.TEXT & ~filters.COMMAND, get_cnpj)],
            REGIME:       [MessageHandler(filters.TEXT & ~filters.COMMAND, get_regime)],
            DATA_INICIO:  [MessageHandler(filters.TEXT & ~filters.COMMAND, get_data)],
            SERVICOS:     [MessageHandler(filters.TEXT & ~filters.COMMAND, get_servicos)],
            FOLHA:        [MessageHandler(filters.TEXT & ~filters.COMMAND, get_folha)],
            HONORARIO:    [MessageHandler(filters.TEXT & ~filters.COMMAND, get_honorario)],
            VENCIMENTO:   [MessageHandler(filters.TEXT & ~filters.COMMAND, get_vencimento)],
            RESPONSAVEL:  [MessageHandler(filters.TEXT & ~filters.COMMAND, get_responsavel)],
            OBS:          [MessageHandler(filters.TEXT & ~filters.COMMAND, get_obs)],
            REPRESENTANTE:[MessageHandler(filters.TEXT & ~filters.COMMAND, get_representante)],
            CPF_REP:      [MessageHandler(filters.TEXT & ~filters.COMMAND, get_cpf_rep)],
            END_EMPRESA:  [MessageHandler(filters.TEXT & ~filters.COMMAND, get_end_empresa)],
            END_REP:      [MessageHandler(filters.TEXT & ~filters.COMMAND, get_end_rep)],
            WHATSAPP:     [MessageHandler(filters.TEXT & ~filters.COMMAND, get_whatsapp)],
            EMAIL:        [MessageHandler(filters.TEXT & ~filters.COMMAND, get_email)],
        },
        fallbacks=[CommandHandler("cancelar", cancelar)],
    )

    conv_boleto = ConversationHandler(
        entry_points=[CommandHandler("boleto", boleto_start)],
        name="boleto",
        persistent=False,
        states={
            BOL_EMPRESA:  [MessageHandler(filters.TEXT & ~filters.COMMAND, bol_empresa)],
            BOL_CNPJ:     [MessageHandler(filters.TEXT & ~filters.COMMAND, bol_cnpj)],
            BOL_DESC:     [MessageHandler(filters.TEXT & ~filters.COMMAND, bol_desc)],
            BOL_VALOR:    [MessageHandler(filters.TEXT & ~filters.COMMAND, bol_valor)],
            BOL_VENC:     [MessageHandler(filters.TEXT & ~filters.COMMAND, bol_venc)],
            BOL_WHATSAPP: [MessageHandler(filters.TEXT & ~filters.COMMAND, bol_whatsapp)],
        },
        fallbacks=[CommandHandler("cancelar", cancelar)],
    )

    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("chatid", chatid))
    app.add_handler(conv)
    app.add_handler(conv_boleto)
    app.add_handler(CommandHandler("meuid", meuid))
    app.add_error_handler(error_handler)
    return app

flask_app = Flask(__name__)
_telegram_app_ref = {"app": None}
_main_event_loop = {"loop": None}


import urllib.parse as _urlparse_helper

URL_BASE_RAILWAY = "https://canvas-contabilidade-production.up.railway.app"


def montar_mensagem_agendor(dados_negocio: dict) -> str:
    org = dados_negocio.get("organization", {}) or {}
    nome_empresa = org.get("name", "Não informado")
    cnpj = org.get("cnpj") or ""
    valor = dados_negocio.get("value", "Não informado")
    whatsapp = (org.get("contact") or {}).get("whatsapp", "") or ""
    titulo = dados_negocio.get("title", "")

    whatsapp_limpo = re.sub(r"\D", "", whatsapp)

    params = {}
    if nome_empresa and nome_empresa != "Não informado":
        params["empresa"] = nome_empresa
    if cnpj:
        params["cnpj"] = cnpj
    if whatsapp_limpo:
        params["whatsapp"] = whatsapp_limpo
    if valor and valor != "Não informado":
        params["honorario"] = str(valor)

    query_string = _urlparse_helper.urlencode(params)
    link_formulario = f"{URL_BASE_RAILWAY}/cadastrar?{query_string}"

    return (
        f"🟢 *Novo negócio ganho no Agendor!*\n\n"
        f"*Negócio:* {titulo}\n"
        f"*Empresa:* {nome_empresa}\n"
        f"*CNPJ:* {cnpj or 'Não informado'}\n"
        f"*Valor:* R$ {valor}\n"
        f"*WhatsApp:* {whatsapp or 'Não informado'}\n\n"
        f"👉 [Clique aqui para completar o cadastro]({link_formulario})\n\n"
        f"(os dados acima já vêm preenchidos no formulário)"
    )


@flask_app.route("/webhook/agendor", methods=["POST"])
def webhook_agendor():
    try:
        payload = request.get_json(force=True)
    except Exception as e:
        return jsonify({"erro": f"JSON invalido: {e}"}), 400

    print("PAYLOAD RECEBIDO DO AGENDOR:", payload)

    dados_negocio = payload.get("data")
    if not dados_negocio:
        return jsonify({"erro": "Campo 'data' nao encontrado no payload"}), 400

    chat_id = os.environ.get("GRUPO_AGENDOR_ID")
    if not chat_id:
        return jsonify({"erro": "GRUPO_AGENDOR_ID nao configurado"}), 500

    mensagem = montar_mensagem_agendor(dados_negocio)

    app_ref = _telegram_app_ref["app"]
    loop = _main_event_loop["loop"]

    if app_ref and loop:
        asyncio.run_coroutine_threadsafe(
            app_ref.bot.send_message(
                chat_id=int(chat_id), text=mensagem, parse_mode="Markdown"
            ),
            loop,
        )

    return jsonify({"status": "ok"}), 200
MAPA_TIPO_FORM = {
    "cnpj": "1",
    "abertura": "3",
    "alteracao": "3",
    "baixa": "3",
    "desenquadramento": "3",
    "alvaras": "2",
}

@flask_app.route("/", methods=["GET"])
@flask_app.route("/cadastrar", methods=["GET"])
def pagina_formulario():
    caminho = Path(__file__).parent / "static_form" / "index.html"
    return caminho.read_text(encoding="utf-8")


@flask_app.route("/cadastrar", methods=["POST"])
def cadastrar_via_formulario():
    log = []

    try:
        payload = request.get_json(force=True)
    except Exception as e:
        return jsonify({"ok": False, "log": [{"tipo": "erro", "msg": f"Dados invalidos: {e}"}]}), 400

    tipo_codigo = MAPA_TIPO_FORM.get(payload.get("tipo", ""), "1")

    d = {
        "tipo": tipo_codigo,
        "empresa": (payload.get("nome") or "").strip().upper(),
        "cnpj": (payload.get("cnpj") or "").strip(),
        "regime": payload.get("regime") or "Não informado",
        "data_inicio": payload.get("inicio_resp") or "",
        "servicos": payload.get("servicos") or payload.get("tipo_servico", ""),
        "folha": payload.get("folha") or "Não",
        "honorario": payload.get("honorario") or "0",
        "vencimento": payload.get("vencimento") or "10",
        "responsavel": payload.get("responsavel") or "",
        "obs": payload.get("observacoes") or "nenhuma",
        "whatsapp": re.sub(r"\D", "", payload.get("telefone") or ""),
        "modalidade": payload.get("modalidade") or "Transferência",
        "representante": payload.get("representante") or "",
        "cpf_rep": payload.get("cpf") or "",
        "end_empresa": payload.get("end_empresa") or "Não informado",
        "end_rep": payload.get("end_rep") or "Não informado",
    }

    if d["data_inicio"] and "-" in d["data_inicio"]:
        partes = d["data_inicio"].split("-")
        if len(partes) == 3:
            d["data_inicio"] = f"{partes[2]}/{partes[1]}/{partes[0]}"

    trello_url = "—"
    card_url = None
    boleto_url = None

    if d["tipo"] == "1":
        try:
            contract_path = fill_contract(d)
            log.append({"tipo": "ok", "msg": "Contrato gerado"})

            if contract_path and contract_path.exists() and D4SIGN_TOKEN:
                try:
                    email_cliente = payload.get("email") or ""
                    uuid_doc = enviar_para_d4sign(
                        contract_path,
                        email_cliente,
                        d["empresa"]
                    )
                    salvar_pendente(uuid_doc, d)
                    log.append({
                        "tipo": "ok",
                        "msg": "Contrato enviado para assinatura via D4Sign"
                    })
                    return jsonify({
                        "ok": True,
                        "log": log,
                        "card_url": None,
                        "boleto_url": None,
                        "whatsapp_url": None,
                        "aguardando_assinatura": True,
                        "mensagem": (
                            "✅ Contrato enviado para assinatura! "
                            "O time operacional será avisado automaticamente "
                            "após o cliente assinar."
                        )
                    }), 200

                except Exception as e:
                    log.append({"tipo": "erro", "msg": f"D4Sign: {e}"})

        except Exception as e:
            log.append({"tipo": "erro", "msg": f"Contrato: {e}"})

    # Se nao for tipo 1, ou se o D4Sign falhou, processa imediatamente
    if d["tipo"] not in SEM_TRELLO:
        try:
            desc = (
                f"Empresa: {d['empresa']}\n"
                f"CNPJ/CPF: {d['cnpj']}\n"
                f"Modalidade: {d['modalidade']}\n"
                f"Regime: {d['regime']}\n"
                f"Inicio: {d['data_inicio']}\n"
                f"Servicos: {d['servicos']}\n"
                f"Folha: {d['folha']}\n"
                f"Honorario: R$ {d['honorario']} - Venc. {d['vencimento']}\n"
                f"Responsavel: {d['responsavel']}\n"
                f"Obs: {d['obs']}"
            )
            nome_cartao = f"{d['empresa']} - {TIPOS[d['tipo']]}"
            trello_url = trello_add_card(d["tipo"], nome_cartao, desc)
            card_url = trello_url
            log.append({"tipo": "ok", "msg": "Cartão criado no Trello"})
        except Exception as e:
            log.append({"tipo": "erro", "msg": f"Trello: {e}"})
    else:
        log.append({"tipo": "info", "msg": "Sem Trello para este serviço"})

    try:
        app_ref = _telegram_app_ref["app"]
        loop = _main_event_loop["loop"]
        if app_ref and loop:
            asyncio.run_coroutine_threadsafe(
                app_ref.bot.send_message(
                    chat_id=GROUP_ID,
                    text=build_group_msg(d, trello_url),
                    parse_mode="Markdown",
                ),
                loop,
            )
            log.append({"tipo": "ok", "msg": "Grupo do Telegram notificado"})
    except Exception as e:
        log.append({"tipo": "erro", "msg": f"Grupo Telegram: {e}"})
        

    if payload.get("forma_pagamento") == "boleto":
        if ASAAS_KEY and ASAAS_KEY != "PREENCHER_COM_CHAVE_ASAAS":
            try:
                valor = float(str(d["honorario"]).replace(".", "").replace(",", "."))
                cid = asaas_customer(d["empresa"], d["cnpj"])
                bol = asaas_boleto(
                    cid, valor, int(d["vencimento"]),
                    f"Honorario {TIPOS[d['tipo']]} — {d['empresa']}",
                )
                boleto_url = bol["url"]
                log.append({"tipo": "ok", "msg": "Boleto gerado no Asaas"})
            except Exception as e:
                log.append({"tipo": "erro", "msg": f"Asaas: {e}"})
        else:
            log.append({"tipo": "erro", "msg": "Asaas: chave não configurada"})
    else:
        log.append({"tipo": "info", "msg": "Pagamento via PIX — sem boleto automático"})

    whatsapp_url = None
    if d["whatsapp"]:
        texto_wpp = (
            f"Olá, {d['empresa'].title()}! 😊\n\n"
            f"Seja bem-vindo(a) à Canvas Contabilidade!\n\n"
        )
        if boleto_url:
            texto_wpp += f"Segue o link do seu boleto:\n{boleto_url}\n\n"
        texto_wpp += "Qualquer dúvida, estamos à disposição! 🟢\nCanvas Contabilidade"
        whatsapp_url = f"https://wa.me/55{d['whatsapp']}?text={urllib.parse.quote(texto_wpp)}"

    return jsonify({
        "ok": True,
        "log": log,
        "card_url": card_url,
        "boleto_url": boleto_url,
        "whatsapp_url": whatsapp_url,
    }), 200

@flask_app.route("/webhook/d4sign", methods=["POST"])
def webhook_d4sign():
    try:
        payload = request.get_json(force=True)
    except Exception:
        payload = {}

    print("WEBHOOK D4SIGN RECEBIDO:", payload)

    uuid_doc = payload.get("uuid_document") or payload.get("uuidDoc")
    tipo_evento = payload.get("type", "")

    if tipo_evento not in ("FINISHED", "SIGNED", "4"):
        return jsonify({"status": "ignorado", "tipo": tipo_evento}), 200

    if not uuid_doc:
        return jsonify({"erro": "uuid_document ausente"}), 400

    d = ler_e_remover_pendente(uuid_doc)
    if not d:
        return jsonify({"erro": "Cadastro nao encontrado para este documento"}), 404

    erros = []
    trello_url = "—"
    boleto_url = "—"

    if d["tipo"] not in SEM_TRELLO:
        try:
            desc = (
                f"Empresa: {d['empresa']}\n"
                f"CNPJ/CPF: {d['cnpj']}\n"
                f"Modalidade: {d.get('modalidade', '')}\n"
                f"Regime: {d['regime']}\n"
                f"Inicio: {d['data_inicio']}\n"
                f"Servicos: {d['servicos']}\n"
                f"Folha: {d['folha']}\n"
                f"Honorario: R$ {d['honorario']} - Venc. {d['vencimento']}\n"
                f"Responsavel: {d['responsavel']}\n"
                f"Obs: {d['obs']}"
            )
            nome_cartao = f"{d['empresa']} - {TIPOS[d['tipo']]}"
            trello_url = trello_add_card(d["tipo"], nome_cartao, desc)
        except Exception as e:
            erros.append(f"Trello: {e}")
    else:
        trello_url = "sem Trello para este servico"

    app_ref = _telegram_app_ref["app"]
    loop = _main_event_loop["loop"]
    if app_ref and loop:
        try:
            asyncio.run_coroutine_threadsafe(
                app_ref.bot.send_message(
                    chat_id=GROUP_ID,
                    text=build_group_msg(d, trello_url),
                    parse_mode="Markdown",
                ),
                loop,
            )
        except Exception as e:
            erros.append(f"Grupo Telegram: {e}")

    if ASAAS_KEY and ASAAS_KEY != "PREENCHER_COM_CHAVE_ASAAS":
        try:
            valor = float(str(d["honorario"]).replace(".", "").replace(",", "."))
            cid = asaas_customer(d["empresa"], d["cnpj"])
            bol = asaas_boleto(
                cid, valor, int(d["vencimento"]),
                f"Honorario {TIPOS[d['tipo']]} — {d['empresa']}",
            )
            boleto_url = bol["url"]
        except Exception as e:
            erros.append(f"Asaas: {e}")

    if d.get("whatsapp") and app_ref and loop:
        try:
            wa_text = (
                f"Olá, {d['empresa'].title()}! 😊\n\n"
                f"Seja bem-vindo(a) à Canvas Contabilidade!\n\n"
                f"Seu contrato foi assinado com sucesso. "
                f"Estamos muito felizes em tê-lo(a) como cliente!\n\n"
            )
            if boleto_url and boleto_url != "—":
                wa_text += f"Segue o link do seu boleto:\n{boleto_url}\n\n"
            wa_text += "Qualquer dúvida, estamos à disposição! 🟢\nCanvas Contabilidade"
            wa_link = f"https://wa.me/55{d['whatsapp']}?text={urllib.parse.quote(wa_text)}"

            asyncio.run_coroutine_threadsafe(
                app_ref.bot.send_message(
                    chat_id=GROUP_ID,
                    text=f"📱 *Link de boas-vindas pronto para enviar:*\n{wa_link}",
                    parse_mode="Markdown",
                    disable_web_page_preview=True,
                ),
                loop,
            )
        except Exception as e:
            erros.append(f"WhatsApp link: {e}")

    return jsonify({"status": "ok", "erros": erros}), 200

def iniciar_flask():
    porta = int(os.environ.get("PORT", 8080))
    flask_app.run(host="0.0.0.0", port=porta, use_reloader=False)


def main():
    while True:
        try:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            app = build_app()

            _telegram_app_ref["app"] = app
            _main_event_loop["loop"] = loop

            flask_thread = threading.Thread(target=iniciar_flask, daemon=True)
            flask_thread.start()

            print("Canvas Bot rodando -- /novo para cadastrar")
            print("Webhook do Agendor disponivel em /webhook/agendor")
            app.run_polling(drop_pending_updates=True)
            break
        except Conflict:
            print("Conflito: outro bot ativo. Aguardando 20s e tentando novamente...")
            time.sleep(20)
        except Exception as e:
            print(f"Erro inesperado: {e}. Reiniciando em 10s...")
            time.sleep(10)

if __name__ == "__main__":
    main()